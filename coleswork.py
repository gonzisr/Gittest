# -*- coding: utf-8 -*-
"""
Created on Sat Oct 15 21:53:00 2016

@author: Israel
"""

import requests
from bs4 import BeautifulSoup
import re
import docx
from docx import Document

class Units():
    
    def _init_(self):
        self.lot = ""
        self.square = ""
        self.address = ""
        self.unit_number = ""
    
def get_forclosure_values(): 
        
    url = "http://realestate.alexcooper.com/upcoming/foreclosures/2016-10-16/0/24/"
    
    r = requests.get(url)
    soup = BeautifulSoup(r.content, "lxml")
    
    links = soup.find_all("a") 
    #link_list = ["<a href='%s'>%s</a>" %(link.get("href"), link.text) for link in links]
    
    ad_number = []
    link_list = [link.get("href") for link in links]  
    
    for lk in link_list:
        if re.findall('\#(\d+)', lk):
             ad_number.append(re.findall('\#(\d+)', lk))
             
    unit_dict = {"ad" : [], "Lot" : [], "Square": []}
    unit_list = []
    
    ad_ref = [x[0] for x in ad_number]
        
    for forclosure_ad in ad_ref:
        
        r = requests.get("http://realestate.alexcooper.com/ajax/foreclosure-ad/%s/" %forclosure_ad)
        soup = BeautifulSoup(r.content, "lxml")
        
        ads = soup.find_all("p")
        
        for units in ads:
            #print units.text
            #if re.findall('Lot (\d+)', units.text): if you want to look for all lots and squares
            if re.findall('Lot \d+ in Square \d+', units.text):
                unit_lot = re.findall('Lot (\d+)', units.text)
                unit_square = re.findall('Square (\d+)', units.text)
                #new_unit = Units()
                #new_unit.lot = re.findall('Lot (\d+)', units.text)
                #new_unit.square = re.findall('Square (\d+)', units.text)
                unit_dict["ad"].append(forclosure_ad)
                unit_dict["Lot"].append(unit_lot) 
                unit_dict["Square"].append(unit_square)
            
            '''
            document = Document() 
            document.add_heading('Lot %s in Square %s'%(new_unit.lot, new_unit.square), 0)
            document.add_paragraph(soup.text, style = "Listnumber")
            
            document.add_page_break()
    
            document.save('demo.docx')'''
        
    for each_unit in unit_list:
        print each_unit.lot
        print each_unit.square
        
    
    
    
    